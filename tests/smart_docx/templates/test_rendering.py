import os
import unittest

from smart_docx.templates.rendering import TemplateRenderer, Field


class TestTemplateRenderer(unittest.TestCase):
    def setUp(self):
        # Create a temporary .docx template for testing.
        self.template_path = "../test_template.docx"
        self.output_path = "../test_output.docx"
        self.fields = [
            Field(id="name", value="Alice"),
            Field(id="date", value="2025-01-18"),
            Field(id="purpose", value="Testing")
        ]

        from docx import Document
        doc = Document()
        doc.add_paragraph("Name: {{ name }}")
        doc.add_paragraph("Date: {{ date }}")
        doc.add_paragraph("Purpose: {{ purpose }}")
        doc.save(self.template_path)

    def tearDown(self):
        # Clean up the temporary .docx files.
        if os.path.exists(self.template_path):
            os.remove(self.template_path)
        if os.path.exists(self.output_path):
            os.remove(self.output_path)

    def test_render_success(self):
        """
        Test successful rendering of the template with valid fields.
        """
        renderer = TemplateRenderer(template_file=self.template_path)
        renderer.render(fields=self.fields, output_path=self.output_path)

        # Assert that the rendered file exists
        self.assertTrue(os.path.exists(self.output_path))

        # Verify the content of the rendered file
        from docx import Document
        rendered_doc = Document(self.output_path)
        paragraphs = [p.text for p in rendered_doc.paragraphs]

        self.assertIn("Name: Alice", paragraphs)
        self.assertIn("Date: 2025-01-18", paragraphs)
        self.assertIn("Purpose: Testing", paragraphs)

    def test_render(self):
        output_path = "./resources/output.docx"
        renderer = TemplateRenderer(template_file="./resources/template.docx")
        renderer.render(fields=[
            Field(id="name", value="Svet"),
            Field(id="items", value=[
                {"name": "Item 1", "description": "Description for item 1"},
                {"name": "Item 2", "description": "Description for item 2"},
                {"name": "Item 3", "description": "Description for item 3"},
            ])
        ],
            output_path=output_path)


if __name__ == "__main__":
    unittest.main()
