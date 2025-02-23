import os
import unittest

from smart_docx.templates.definitions import TemplateDefinition, FieldDefinition, SourceType
from smart_docx.templates.validations import validate_template_variables


class TestTemplateValidator(unittest.TestCase):
    def setUp(self):
        self.template_path = "test_template.docx"
        from docx import Document
        doc = Document()
        doc.add_paragraph("Name: {{ name }}")
        doc.add_paragraph("Date: {{ date }}")
        doc.add_paragraph("Purpose: {{r purpose }}")
        doc.save(self.template_path)

    def tearDown(self):
        if os.path.exists(self.template_path):
            os.remove(self.template_path)

    def test_validate_template_variables(self):
        # valid template definition
        template_def = TemplateDefinition(
            name="test",
            file=self.template_path,
            instructions="instructions ...",
            description="description ...",
            fields=[
                FieldDefinition(
                    id="name",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Get name of the person"
                ),
                FieldDefinition(
                    id="date",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Get date of birth"
                ),
                FieldDefinition(
                    id="purpose",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Purpose of event"
                ),
            ]
        )

        validate_template_variables(template_def)

    def test_validate_invalid_variables(self):
        # missing field definition
        invalid_template = TemplateDefinition(
            name="test",
            file=self.template_path,
            instructions="instructions ...",
            description="description ...",
            fields=[
                FieldDefinition(
                    id="name",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Get name of the person"
                ),
                FieldDefinition(
                    id="date",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Get date of birth"
                ),
            ]
        )

        with self.assertRaises(ValueError) as context:
            validate_template_variables(invalid_template)
        self.assertEqual("Missing defined fields, which are present in template: purpose", str(context.exception))

        # circular dependencies
        circular_dependencies = TemplateDefinition(
            file=self.template_path,
            name="circular_dependencies",
            instructions="instructions ...",
            description="description ...",
            fields=[
                FieldDefinition(
                    id="name",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Get name of the person {{ purpose }}"
                ),
                FieldDefinition(
                    id="date",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Get date of birth {{ name }}"
                ),
                FieldDefinition(
                    id="purpose",
                    source=SourceType.AUTO,
                    value={"type": "string"},
                    instructions="Purpose of event {{ date }}"
                ),
            ]
        )

        with self.assertRaises(ValueError) as context:
            validate_template_variables(circular_dependencies)
        self.assertEqual("Circular dependency between fields", str(context.exception))


if __name__ == "__main__":
    unittest.main()
